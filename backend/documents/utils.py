import os
import uuid
import subprocess
import shutil
import tempfile
import logging
from datetime import datetime
from django.conf import settings
from django.core.files.storage import default_storage
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

# WeasyPrint - gestion des erreurs sur Windows
WEASYPRINT_AVAILABLE = False
HTML = None
CSS = None
FontConfiguration = None

try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
    logger.info("✅ WeasyPrint chargé avec succès")
except ImportError as e:
    logger.warning(f"⚠️ WeasyPrint non disponible (ImportError): {e}")
    WEASYPRINT_AVAILABLE = False
except OSError as e:
    logger.warning(f"⚠️ WeasyPrint non disponible (OSError - dépendances manquantes): {e}")
    WEASYPRINT_AVAILABLE = False
except Exception as e:
    logger.warning(f"⚠️ WeasyPrint non disponible (erreur inattendue): {e}")
    WEASYPRINT_AVAILABLE = False


class DocumentGenerator:
    """
    Classe principale pour la génération de documents PDF à partir de templates Word.
    """
    
    def __init__(self):
        self.templates_dir = os.path.join(settings.BASE_DIR, 'templates', 'word')
        self.generated_dir = os.path.join(settings.MEDIA_ROOT, 'generated')
        self.ensure_directories()
    
    def ensure_directories(self):
        """S'assurer que les répertoires existent."""
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.generated_dir, exist_ok=True)
    
    def generate_docx_from_template(self, template_name, data, output_filename=None):
        """
        Génère un document Word (.docx) à partir d'un template.
        """
        try:
            # Vérifier que le template existe
            if not self.validate_template(template_name):
                return {
                    'success': False,
                    'error': f'Template {template_name} non trouvé'
                }
            
            # Chemin du template
            template_path = os.path.join(self.templates_dir, template_name)
            
            # Nom du fichier de sortie
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"{template_name}_{timestamp}.docx"
            
            # Chemin de sortie
            output_path = os.path.join(self.generated_dir, output_filename)
            
            # Charger le template
            template = DocxTemplate(template_path)
            
            # Rendre le template avec les données
            template.render(data)
            
            # Sauvegarder le document généré
            template.save(output_path)
            
            # Vérifier que le fichier a été créé
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"Document Word généré: {output_path} ({file_size} octets)")
                
                return {
                    'success': True,
                    'file_path': output_path,
                    'filename': output_filename,
                    'file_size': file_size
                }
            else:
                return {
                    'success': False,
                    'error': 'Erreur lors de la sauvegarde du document'
                }
                
        except Exception as e:
            logger.error(f"Erreur lors de la génération du document Word: {e}")
            return {
                'success': False,
                'error': f'Erreur lors de la génération: {str(e)}'
            }

    def generate_pdf_from_word_template(self, template_name, data, output_filename=None):
        """
        Génère un PDF à partir d'un template Word (.docx).
        Sur Windows, génère directement un fichier DOCX si la conversion PDF échoue.
        
        Args:
            template_name (str): Nom du fichier template (ex: 'fiche_projet_marketing.docx')
            data (dict): Données à injecter dans le template
            output_filename (str): Nom du fichier de sortie (optionnel)
        
        Returns:
            dict: Informations sur le fichier généré
        """
        try:
            # Chemin du template
            template_path = os.path.join(self.templates_dir, template_name)
            
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template non trouvé: {template_path}")
            
            # Générer un nom de fichier unique si non fourni
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                unique_id = str(uuid.uuid4())[:8]
                output_filename = f"{template_name.replace('.docx', '')}_{timestamp}_{unique_id}.pdf"
            
            # Chemin de sortie
            output_path = os.path.join(self.generated_dir, output_filename)
            
            # Charger le template Word
            doc = DocxTemplate(template_path)
            
            # Rendre le template avec les données
            doc.render(data)
            
            # Sauvegarder temporairement le document Word modifié
            temp_docx_path = os.path.join(tempfile.gettempdir(), f"temp_{uuid.uuid4()}.docx")
            doc.save(temp_docx_path)
            
            try:
                # Essayer de convertir en PDF
                pdf_path = self.convert_docx_to_pdf(temp_docx_path, output_path)
                final_path = pdf_path
                final_filename = output_filename
            except Exception as pdf_error:
                # Fallback: sauvegarder le DOCX directement
                logger.warning(f"Conversion PDF échouée: {pdf_error}. Sauvegarde en DOCX.")
                docx_filename = output_filename.replace('.pdf', '.docx')
                docx_path = os.path.join(self.generated_dir, docx_filename)
                shutil.copy2(temp_docx_path, docx_path)
                final_path = docx_path
                final_filename = docx_filename
            
            # Nettoyer le fichier temporaire
            if os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
            
            # Calculer la taille du fichier
            file_size = os.path.getsize(final_path) if os.path.exists(final_path) else 0
            
            return {
                'success': True,
                'file_path': final_path,
                'filename': final_filename,
                'file_size': file_size,
                'template_used': template_name,
                'generated_at': datetime.now().isoformat(),
                'format': 'pdf' if final_filename.endswith('.pdf') else 'docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'template_used': template_name
            }
    
    def convert_docx_to_pdf(self, docx_path, pdf_path):
        """
        Convertit un fichier DOCX en PDF.
        Utilise LibreOffice en priorité, puis fallback vers HTML/WeasyPrint.
        """
        try:
            # Méthode 1: LibreOffice (recommandée pour la qualité)
            if self._convert_with_libreoffice(docx_path, pdf_path):
                logger.info(f"PDF généré avec LibreOffice: {pdf_path}")
                return pdf_path
            
            # Méthode 2: WeasyPrint via HTML
            logger.info("LibreOffice non disponible, utilisation de WeasyPrint")
            from docx import Document
            
            # Lire le document Word
            doc = Document(docx_path)
            
            # Convertir en HTML
            html_content = self.docx_to_html(doc)
            
            # Générer le PDF à partir du HTML
            self.html_to_pdf(html_content, pdf_path)
            
            return pdf_path
            
        except Exception as e:
            # Méthode de fallback: copier le fichier DOCX
            fallback_path = pdf_path.replace('.pdf', '.docx')
            shutil.copy2(docx_path, fallback_path)
            logger.error(f"Erreur de conversion PDF: {e}")
            raise Exception(f"Erreur de conversion PDF: {e}. Fichier DOCX sauvegardé: {fallback_path}")
    
    def _convert_with_libreoffice(self, docx_path, pdf_path):
        """
        Convertit DOCX en PDF avec LibreOffice.
        """
        try:
            # Vérifier si LibreOffice est disponible
            result = subprocess.run(['libreoffice', '--version'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode != 0:
                return False
            
            # Créer le répertoire de sortie
            output_dir = os.path.dirname(pdf_path)
            os.makedirs(output_dir, exist_ok=True)
            
            # Commande LibreOffice pour convertir
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ]
            
            # Exécuter la conversion
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0:
                # LibreOffice génère le PDF avec le même nom que le DOCX
                expected_pdf = os.path.join(output_dir, 
                    os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
                
                if os.path.exists(expected_pdf):
                    # Renommer vers le nom souhaité
                    if expected_pdf != pdf_path:
                        shutil.move(expected_pdf, pdf_path)
                    return True
            
            return False
            
        except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
            logger.warning(f"LibreOffice non disponible ou erreur: {e}")
            return False
    
    def docx_to_html(self, doc):
        """
        Convertit un document Word en HTML.
        """
        html_parts = []
        html_parts.append('<html><head><meta charset="utf-8">')
        html_parts.append('<style>')
        html_parts.append('body { font-family: Arial, sans-serif; margin: 40px; }')
        html_parts.append('h1, h2, h3 { color: #333; }')
        html_parts.append('table { border-collapse: collapse; width: 100%; }')
        html_parts.append('td, th { border: 1px solid #ddd; padding: 8px; text-align: left; }')
        html_parts.append('th { background-color: #f2f2f2; }')
        html_parts.append('</style></head><body>')
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Déterminer le style selon le niveau de titre
                if paragraph.style.name.startswith('Heading 1'):
                    html_parts.append(f'<h1>{paragraph.text}</h1>')
                elif paragraph.style.name.startswith('Heading 2'):
                    html_parts.append(f'<h2>{paragraph.text}</h2>')
                elif paragraph.style.name.startswith('Heading 3'):
                    html_parts.append(f'<h3>{paragraph.text}</h3>')
                else:
                    html_parts.append(f'<p>{paragraph.text}</p>')
        
        # Traiter les tableaux
        for table in doc.tables:
            html_parts.append('<table>')
            for i, row in enumerate(table.rows):
                html_parts.append('<tr>')
                for cell in row.cells:
                    tag = 'th' if i == 0 else 'td'
                    html_parts.append(f'<{tag}>{cell.text}</{tag}>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('</body></html>')
        return ''.join(html_parts)
    
    def html_to_pdf(self, html_content, pdf_path):
        """
        Convertit du HTML en PDF.
        """
        if not WEASYPRINT_AVAILABLE or HTML is None or CSS is None or FontConfiguration is None:
            # Fallback: utiliser ReportLab pour générer un PDF simple
            logger.info("WeasyPrint non disponible, utilisation de ReportLab comme fallback")
            return self._generate_pdf_with_reportlab(html_content, pdf_path)
        
        try:
            font_config = FontConfiguration()
            
            html_doc = HTML(string=html_content)
            html_doc.write_pdf(
                pdf_path,
                font_config=font_config,
                stylesheets=[CSS(string='''
                    @page {
                        size: A4;
                        margin: 2cm;
                    }
                    body {
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                    }
                ''')]
            )
            logger.info(f"PDF généré avec WeasyPrint: {pdf_path}")
        except Exception as e:
            logger.warning(f"Erreur WeasyPrint, fallback vers ReportLab: {e}")
            return self._generate_pdf_with_reportlab(html_content, pdf_path)
    
    def _generate_pdf_with_reportlab(self, html_content, pdf_path):
        """
        Génère un PDF simple avec ReportLab en fallback.
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            import re
            
            # Nettoyer le HTML pour extraire le texte
            clean_text = re.sub(r'<[^>]+>', '', html_content)
            clean_text = clean_text.replace('&nbsp;', ' ').replace('&amp;', '&')
            
            # Créer le PDF
            doc = SimpleDocTemplate(pdf_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Style personnalisé
            custom_style = ParagraphStyle(
                'CustomStyle',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leftIndent=0,
                rightIndent=0
            )
            
            # Diviser le texte en paragraphes
            paragraphs = clean_text.split('\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), custom_style))
                    story.append(Spacer(1, 6))
            
            # Construire le PDF
            doc.build(story)
            logger.info(f"PDF généré avec ReportLab: {pdf_path}")
            return True
            
        except Exception as e:
            logger.error(f"Erreur ReportLab: {e}")
            # Dernier recours: sauvegarder en HTML
            html_path = pdf_path.replace('.pdf', '.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            raise Exception(f"Impossible de générer le PDF. HTML sauvegardé: {html_path}")
    
    def get_available_templates(self):
        """
        Retourne la liste des templates disponibles.
        """
        templates = []
        if os.path.exists(self.templates_dir):
            for filename in os.listdir(self.templates_dir):
                if filename.endswith('.docx'):
                    templates.append({
                        'name': filename,
                        'path': os.path.join(self.templates_dir, filename),
                        'size': os.path.getsize(os.path.join(self.templates_dir, filename))
                    })
        return templates
    
    def validate_template(self, template_name):
        """
        Valide qu'un template existe et est accessible.
        """
        template_path = os.path.join(self.templates_dir, template_name)
        return os.path.exists(template_path) and os.access(template_path, os.R_OK)
    
    def get_template_info(self, template_name):
        """
        Retourne les informations détaillées d'un template.
        """
        template_path = os.path.join(self.templates_dir, template_name)
        if not os.path.exists(template_path):
            return None
        
        stat = os.stat(template_path)
        return {
            'name': template_name,
            'path': template_path,
            'size': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'exists': True
        }
    
    def extract_template_variables(self, template_name):
        """
        Extrait les variables d'un template Word.
        """
        try:
            template_path = os.path.join(self.templates_dir, template_name)
            if not os.path.exists(template_path):
                return []
            
            doc = DocxTemplate(template_path)
            # Récupérer les variables du template
            variables = []
            
            # Analyser les paragraphes pour trouver les variables
            for paragraph in doc.template.docx.paragraphs:
                text = paragraph.text
                # Rechercher les variables au format {{variable}}
                import re
                matches = re.findall(r'\{\{(\w+)\}\}', text)
                variables.extend(matches)
            
            # Analyser les tableaux
            for table in doc.template.docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            matches = re.findall(r'\{\{(\w+)\}\}', text)
                            variables.extend(matches)
            
            return list(set(variables))  # Supprimer les doublons
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction des variables du template {template_name}: {e}")
            return []


class TemplateManager:
    """
    Gestionnaire pour les templates de documents.
    """
    
    def __init__(self):
        self.generator = DocumentGenerator()
    
    def create_template_mapping(self):
        """
        Crée le mapping entre les types de documents et leurs templates.
        """
        return {
            'fiche_projet_marketing': 'fiche_projet_marketing.docx',
            'fiche_plan_projet': 'fiche_plan_projet.docx',
            'fiche_analyse_offre': 'fiche_analyse_offre.docx',
            'fiche_test': 'fiche_test.docx',
            'fiche_implementation_technique': 'fiche_implementation_technique.docx',
            'fiche_suppression_offre': 'fiche_suppression_offre.docx',
            'specifications_marketing_offre': 'specifications_marketing_offre.docx',
            'ordre_travaux': 'ordre_travaux.docx',
            'fiche_etude_si': 'fiche_etude_si.docx',
            'fiche_etude_technique': 'fiche_etude_technique.docx',
            'fiche_etude_financiere': 'fiche_etude_financiere.docx',
            'fiche_specifications_marketing': 'fiche_specifications_marketing.docx',
            'fiche_implementation': 'fiche_implementation.docx',
            'fiche_recette_uat': 'fiche_recette_uat.docx',
            'fiche_lancement_commercial': 'fiche_lancement_commercial.docx',
            'fiche_projet_complete': 'fiche_projet_complete.docx',
            'contrat': 'contrat.docx',
            'devis': 'devis.docx',
            'facture': 'facture.docx',
            'fiche_suppression': 'fiche_suppression.docx',
            'fiche_bilan_3_mois': 'fiche_bilan_3_mois.docx',
            'fiche_bilan_6_mois': 'fiche_bilan_6_mois.docx',
        }
    
    def get_template_for_document_type(self, document_type):
        """
        Retourne le nom du template pour un type de document donné.
        """
        mapping = self.create_template_mapping()
        return mapping.get(document_type)
    
    def generate_document(self, document_type, data, output_filename=None):
        """
        Génère un document PDF pour un type donné.
        """
        template_name = self.get_template_for_document_type(document_type)
        
        if not template_name:
            return {
                'success': False,
                'error': f'Aucun template trouvé pour le type de document: {document_type}'
            }
        
        if not self.generator.validate_template(template_name):
            return {
                'success': False,
                'error': f'Template non trouvé ou inaccessible: {template_name}'
            }
        
        return self.generator.generate_pdf_from_word_template(
            template_name, data, output_filename
        )
    
    def get_all_templates_info(self):
        """
        Retourne les informations de tous les templates disponibles.
        """
        templates = []
        mapping = self.create_template_mapping()
        
        for doc_type, template_name in mapping.items():
            info = self.generator.get_template_info(template_name)
            if info:
                info['document_type'] = doc_type
                info['variables'] = self.generator.extract_template_variables(template_name)
                templates.append(info)
        
        return templates
    
    def validate_template_for_document_type(self, document_type):
        """
        Valide qu'un template existe pour un type de document.
        """
        template_name = self.get_template_for_document_type(document_type)
        if not template_name:
            return False
        return self.generator.validate_template(template_name)
    
    def get_template_variables(self, document_type):
        """
        Retourne les variables d'un template pour un type de document.
        """
        template_name = self.get_template_for_document_type(document_type)
        if not template_name:
            return []
        return self.generator.extract_template_variables(template_name)
